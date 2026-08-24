from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "guides" / "deskmate-three-end-development-guide-2026-08-24.md"
OUTPUT = ROOT / "docs" / "guides" / "DeskMate-three-end-development-guide-2026-08-24.docx"

BLUE = "2E74B5"
DEEP_BLUE = "1F4D78"
INK = "23313F"
MUTED = "667788"
PALE_BLUE = "E8EEF5"
PALE_CYAN = "EAF6F8"
PALE_ORANGE = "FFF3E8"
PALE_RED = "FDECEC"
PALE_GREEN = "EAF5EA"
GRID = "CAD5DF"
WHITE = "FFFFFF"

CONTENT_WIDTH_DXA = 9360


def set_east_asia(run, name="Microsoft YaHei"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)


def set_run(run, *, size=None, bold=None, color=None, font="Microsoft YaHei", italic=None):
    set_east_asia(run, font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    assert sum(widths) == CONTENT_WIDTH_DXA
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def create_decimal_abstract(document):
    numbering = document.part.numbering_part.element
    existing = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(existing, default=-1) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "%1.")
    level.append(level_text)
    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    level.append(level_jc)
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    level.append(p_pr)
    abstract.append(level)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)
    return abstract_id


def create_number_instance(document, abstract_id):
    numbering = document.part.numbering_part.element
    existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    override.append(start_override)
    num.append(override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(num_id_node)
    p_pr.append(num_pr)


def set_keep_together(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def paragraph_border(paragraph, side="start", color=BLUE, size="18", space="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), size)
    edge.set(qn("w:space"), space)
    edge.set(qn("w:color"), color)
    p_bdr.append(edge)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run(run, size=8.5, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    tail = paragraph.add_run(" 页")
    set_run(tail, size=8.5, color=MUTED)


def add_inline_runs(paragraph, text, *, base_size=11, base_color=INK, bold=False):
    text = text.rstrip()
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run(run, size=base_size, color=base_color, bold=bold)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run(run, size=max(base_size - 0.5, 8), color=DEEP_BLUE, font="Consolas")
            run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        else:
            run = paragraph.add_run(token[2:-2])
            set_run(run, size=base_size, color=base_color, bold=True)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run(run, size=base_size, color=base_color, bold=bold)


def normalize_markdown_lines(text):
    lines = text.replace("\r\n", "\n").split("\n")
    normalized = []
    for line in lines:
        if line.endswith("  "):
            line = line[:-2]
        normalized.append(line)
    return normalized


def parse_table(lines, start):
    rows = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[index].strip().strip("|").split("|")]
        rows.append(cells)
        index += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell) for cell in rows[1]):
        rows.pop(1)
    return rows, index


def table_widths(rows):
    columns = len(rows[0])
    scores = []
    for col in range(columns):
        longest = max(len(re.sub(r"[`*]", "", row[col])) if col < len(row) else 0 for row in rows)
        scores.append(max(9, min(longest, 34)))
    minimum = 1200 if columns <= 5 else 900
    widths = [max(minimum, round(CONTENT_WIDTH_DXA * score / sum(scores))) for score in scores]
    overflow = sum(widths) - CONTENT_WIDTH_DXA
    if overflow > 0:
        adjustable = [i for i, width in enumerate(widths) if width > minimum]
        while overflow > 0 and adjustable:
            share = max(1, overflow // len(adjustable))
            next_adjustable = []
            for i in adjustable:
                delta = min(share, widths[i] - minimum, overflow)
                widths[i] -= delta
                overflow -= delta
                if widths[i] > minimum:
                    next_adjustable.append(i)
                if overflow == 0:
                    break
            adjustable = next_adjustable
    elif overflow < 0:
        widths[-1] += -overflow
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_markdown_table(document, rows):
    if not rows:
        return
    columns = len(rows[0])
    table = document.add_table(rows=len(rows), cols=columns)
    table.style = "Table Grid"
    widths = table_widths(rows)
    set_table_geometry(table, widths)
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        set_keep_together(row)
        if row_index == 0:
            set_repeat_table_header(row)
        for col_index in range(columns):
            cell = row.cells[col_index]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            value = values[col_index] if col_index < len(values) else ""
            add_inline_runs(
                paragraph,
                value,
                base_size=9.2,
                base_color=INK if row_index else DEEP_BLUE,
                bold=row_index == 0,
            )
            if row_index == 0:
                set_cell_shading(cell, PALE_BLUE)
            elif row_index % 2 == 0:
                set_cell_shading(cell, "F7F9FB")
    after = document.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_callout(document, text, fill=PALE_CYAN, color=DEEP_BLUE):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.16)
    paragraph.paragraph_format.right_indent = Inches(0.10)
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(9)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph_shading(paragraph, fill)
    paragraph_border(paragraph, color=color)
    add_inline_runs(paragraph, text, base_size=10.5, base_color=color, bold=False)


def add_code_block(document, lines):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.10)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.05
    paragraph_shading(paragraph, "F3F5F7")
    paragraph_border(paragraph, color="7A8997", size="10", space="6")
    run = paragraph.add_run("\n".join(lines))
    set_run(run, size=9.2, color=INK, font="Consolas")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_markdown_image(document, alt_text, source_path):
    image_path = (SOURCE.parent / source_path).resolve()
    if not image_path.is_file():
        raise FileNotFoundError(f"Markdown image not found: {image_path}")

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run()
    picture = run.add_picture(str(image_path), width=Inches(5.9))
    picture._inline.docPr.set("descr", alt_text or image_path.stem)
    picture._inline.docPr.set("title", alt_text or image_path.stem)

    if alt_text:
        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(10)
        caption.paragraph_format.keep_with_next = False
        run = caption.add_run(f"图 1 · {alt_text}")
        set_run(run, size=9, color=MUTED)


def add_title_paragraph(document, text, level):
    style_name = {2: "Heading 1", 3: "Heading 2"}[level]
    paragraph = document.add_paragraph(style=style_name)
    paragraph.paragraph_format.keep_with_next = True
    add_inline_runs(paragraph, text, base_size=16 if level == 2 else 12.5, base_color=BLUE if level == 2 else DEEP_BLUE, bold=True)
    return paragraph


def build_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ("Body Text", "List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)

    h1 = styles["Heading 1"]
    h1.font.name = "Microsoft YaHei"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(BLUE)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.keep_together = True

    h2 = styles["Heading 2"]
    h2.font.name = "Microsoft YaHei"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    h2.font.size = Pt(12.5)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(DEEP_BLUE)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.keep_together = True


def add_cover(document, source_text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(52)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run("DESKMATE  /  DEVELOPMENT GUIDE")
    set_run(run, size=10, bold=True, color=BLUE)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(10)
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("三端软硬件\n开发指导书")
    set_run(run, size=31, bold=True, color=INK)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(22)
    run = subtitle.add_run("Windows 软件 × EasyInput 总控 × 小智云台")
    set_run(run, size=15, color=DEEP_BLUE)

    line = document.add_paragraph()
    line.paragraph_format.space_after = Pt(20)
    paragraph_border(line, side="bottom", color=BLUE, size="28", space="4")

    meta = document.add_paragraph()
    meta.paragraph_format.space_after = Pt(18)
    add_inline_runs(meta, "专用最新版  ·  2026-08-24  ·  面向第一次进行双 ESP32 与 Windows 集成的开发者", base_size=10.5, base_color=MUTED)

    add_callout(
        document,
        "主原则：相似功能只组成一个小功能包。每完成一包，立刻做针对性测试、三端连通测试和既有锁定功能全量回归；全部通过后记录并锁定，再开始下一包。",
        fill=PALE_CYAN,
        color=DEEP_BLUE,
    )

    purpose = document.add_paragraph()
    purpose.paragraph_format.space_before = Pt(20)
    purpose.paragraph_format.space_after = Pt(6)
    run = purpose.add_run("本指导书解决三件事")
    set_run(run, size=12, bold=True, color=INK)
    for item in (
        "两块 ESP32 板如何安全接线、如何建立第一条可靠控制链路；",
        "软件、EasyInput 固件和小智固件按什么顺序开发与验收；",
        "如何逐步加入本地语音陪伴、长期记忆、说话人识别、表情与安全动作。",
    ):
        p = document.add_paragraph(style="List Bullet")
        add_inline_runs(p, item, base_size=10.5)

    warning = document.add_paragraph()
    warning.paragraph_format.space_before = Pt(24)
    warning.paragraph_format.space_after = Pt(0)
    run = warning.add_run("硬件安全说明")
    set_run(run, size=9.5, bold=True, color="9B3A2A")
    warning2 = document.add_paragraph()
    warning2.paragraph_format.space_after = Pt(0)
    add_inline_runs(
        warning2,
        "本文件不会授权烧录或带电接线。任何真机写入前都必须重新确认设备身份、供电和目标固件；J4 3V3 不接小智。",
        base_size=9.5,
        base_color="9B3A2A",
    )
    document.add_page_break()


def add_navigation(document):
    heading = document.add_paragraph(style="Heading 1")
    add_inline_runs(heading, "阅读地图", base_size=18, base_color=BLUE, bold=True)
    intro = document.add_paragraph()
    add_inline_runs(intro, "第一次使用时按阶段顺序执行；遇到问题时按下表回到对应合同和验收门。", base_size=11)

    rows = [
        ["路线", "章节", "要解决的问题"],
        ["基础与边界", "1–4", "三端分工、真实基线、安全接线、总体架构"],
        ["通信合同", "5", "DeskMate Link v1 的 framing、消息、确认、超时与错误"],
        ["产品体验", "6–9", "两种语音模式、长期记忆、OLED、表情与动作安全"],
        ["实施路径", "10–13", "功能包门禁、阶段计划、三端验收和烧录规则"],
        ["协作与证据", "14–17", "用户操作、立即顺序、资料来源与最终提醒"],
    ]
    add_markdown_table(document, rows)

    loop = document.add_paragraph()
    loop.paragraph_format.space_before = Pt(16)
    loop.paragraph_format.space_after = Pt(8)
    run = loop.add_run("每个功能包的唯一允许流程")
    set_run(run, size=13, bold=True, color=DEEP_BLUE)
    add_code_block(
        document,
        [
            "定义单一目标 → 只改一个功能包 → 自动/模拟测试",
            "        ↓",
            "三端连通测试 → 既有锁定功能全量回归 → 记录证据并锁定",
            "        ↓",
            "只有全部通过，才进入下一个功能包",
        ],
    )
    add_callout(
        document,
        "出现失败时，不继续叠加新功能。先判断是软件、EasyInput、DeskMate Link 还是小智执行端；修复后重新跑当前功能包和所有已锁定回归项。",
        fill=PALE_ORANGE,
        color="9B5C23",
    )
    document.add_page_break()


def add_body_from_markdown(document, text):
    lines = normalize_markdown_lines(text)
    index = 0
    in_code = False
    code_lines = []
    started = False
    page_break_sections = {"3.", "7.", "10.", "11.", "12.", "16."}
    decimal_abstract_id = create_decimal_abstract(document)
    active_number_id = None

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue
        if not started:
            if stripped.startswith("## "):
                started = True
            else:
                index += 1
                continue
        if not stripped:
            active_number_id = None
            index += 1
            continue
        if stripped.startswith("## "):
            active_number_id = None
            text_value = stripped[3:]
            if any(text_value.startswith(prefix) for prefix in page_break_sections):
                document.add_page_break()
            add_title_paragraph(document, text_value, 2)
            index += 1
            continue
        if stripped.startswith("### "):
            active_number_id = None
            text_value = stripped[4:]
            if text_value.startswith("7.3 "):
                document.add_page_break()
            add_title_paragraph(document, text_value, 3)
            index += 1
            continue
        image_match = re.fullmatch(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if image_match:
            active_number_id = None
            add_markdown_image(document, image_match.group(1), image_match.group(2))
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines):
            active_number_id = None
            rows, index = parse_table(lines, index)
            add_markdown_table(document, rows)
            continue
        if stripped.startswith("> "):
            active_number_id = None
            add_callout(document, stripped[2:])
            index += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            active_number_id = None
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, bullet.group(1), base_size=10.7)
            index += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            if active_number_id is None:
                active_number_id = create_number_instance(document, decimal_abstract_id)
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.25
            apply_numbering(paragraph, active_number_id)
            add_inline_runs(paragraph, numbered.group(1), base_size=10.7)
            index += 1
            continue

        active_number_id = None
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.widow_control = True
        add_inline_runs(paragraph, stripped, base_size=10.7)
        index += 1


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_p.add_run("DESKMATE  /  THREE-END DEVELOPMENT")
    set_run(run, size=8, bold=True, color=MUTED)
    paragraph_border(header_p, side="bottom", color="D6DEE6", size="6", space="5")

    footer = section.footer
    footer_p = footer.paragraphs[0]
    add_page_number(footer_p)

    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""
    first_footer = section.first_page_footer
    first_footer.paragraphs[0].text = ""

    settings = document.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    props = document.core_properties
    props.title = "DeskMate 三端软硬件开发指导书"
    props.subject = "Windows 软件、EasyInput 总控固件与小智云台固件的分阶段集成路线"
    props.author = "DeskMate Project"
    props.keywords = "DeskMate, EasyInput, ESP32-S3, Xiaozhi, UART, firmware, local memory"
    props.comments = "Generated from the repository development guide source."


def main():
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    source_text = SOURCE.read_text(encoding="utf-8")
    document = Document()
    configure_document(document)
    build_styles(document)
    add_cover(document, source_text)
    add_navigation(document)
    add_markdown_image(
        document,
        "DeskMate V1 硬件基线高密度信息图",
        "../assets/hardware/deskmate-v1-hardware-baseline-infographic.png",
    )
    document.add_page_break()
    add_body_from_markdown(document, source_text)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    try:
        main()
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        raise
